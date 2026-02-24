"""DOCX text extractor using python-docx."""

import logging

from localsearch.extractors.base import BaseExtractor, ExtractionError, ExtractionResult

logger = logging.getLogger(__name__)


class DocxExtractor(BaseExtractor):
    """Extracts text from DOCX files using python-docx."""

    def extract(self, file_path: str) -> ExtractionResult:
        try:
            from docx import Document
        except ImportError:
            raise ExtractionError("python-docx not installed. Install with: pip install python-docx")

        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Also extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_texts.append(row_text)

            all_text = "\n".join(paragraphs)
            if table_texts:
                all_text += "\n\n" + "\n".join(table_texts)

            if not all_text.strip():
                raise ExtractionError(f"No text extracted from DOCX: {file_path}")

            return ExtractionResult(
                text=all_text,
                metadata={
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables),
                },
            )
        except ExtractionError:
            raise
        except Exception as e:
            raise ExtractionError(f"Failed to extract DOCX {file_path}: {e}") from e

    def supported_extensions(self) -> list[str]:
        return [".docx"]
